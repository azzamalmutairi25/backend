#!/usr/bin/env python3
# ════════════════════════════════════════════════════════════
#  بناء نسخة Word من دليل المستخدم.
#
#  النسخة المسلَّمة تُقرأ وتُعلَّق عليها وتُدمَج في محاضر الجهة، وهذا لا يقع على
#  HTML ولا على PDF. فيُولَّد ملف .docx حقيقي من المصدر نفسه (USER_GUIDE.md)،
#  لا ملف HTML بامتداد Word: عناوين Word حقيقية تظهر في لوحة التنقّل، وجداول
#  حقيقية تُحرَّر، وصور مضمَّنة، واتجاه من اليمين إلى اليسار على مستوى المقطع
#  والفقرة والجدول — لا محاذاةً يمينيةً تنكسر عند أول تعديل.
#
#  الاستعمال:
#      python3 docs/build-guide-docx.py
# ════════════════════════════════════════════════════════════
import os
import re
import struct
import subprocess
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, 'USER_GUIDE.md')
OUT = os.path.join(HERE, 'USER_GUIDE.docx')
IMG = os.path.join(HERE, 'guide-images')
CACHE = os.path.join(HERE, '.guide-build-cache')

FONT = 'Arial'
MONO = 'Consolas'
ACCENT = RGBColor(0x0F, 0x6B, 0x45)
MUTED = RGBColor(0x6B, 0x7F, 0x74)
MAX_W = 6.4          # بوصة — عرض النصّ في A4 بهوامش ١٫٤سم
MAX_H = 7.6          # بوصة — كيلا تبتلع لقطةٌ طويلة صفحةً كاملة


# ── أدوات XML خام: python-docx لا يُعرّض ثنائية الاتجاه ──────
def _set(el, tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    el.append(e)
    return e


def rtl_paragraph(p):
    """اتجاه الفقرة من اليمين — لا محاذاةً فحسب."""
    pPr = p._p.get_or_add_pPr()
    _set(pPr, 'w:bidi')
    return p


def rtl_run(run, on=True):
    rPr = run._r.get_or_add_rPr()
    _set(rPr, 'w:rtl', **{'w:val': '1' if on else '0'})
    _set(rPr, 'w:lang', **{'w:bidi': 'ar-SA'})


def shade(el, fill):
    pPr = el.get_or_add_pPr() if hasattr(el, 'get_or_add_pPr') else el
    _set(pPr, 'w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': fill})


def png_size(path):
    with open(path, 'rb') as f:
        head = f.read(26)
    if head[:8] == b'\x89PNG\r\n\x1a\n':
        w, h = struct.unpack('>II', head[16:24])
        return w, h
    return None


def jpeg_size(path):
    with open(path, 'rb') as f:
        data = f.read()
    i = 2
    while i < len(data) - 9:
        if data[i] != 0xFF:
            i += 1
            continue
        marker = data[i + 1]
        if marker in (0xC0, 0xC1, 0xC2, 0xC3):
            h, w = struct.unpack('>HH', data[i + 5:i + 9])
            return w, h
        seg = struct.unpack('>H', data[i + 2:i + 4])[0]
        i += 2 + seg
    return None


def picture_path(name):
    """يُفضَّل المصغَّر المضغوط إن وُجد — يردّ حجم الملف إلى الثلث."""
    jpg = os.path.join(CACHE, name + '.jpg')
    png = os.path.join(IMG, name + '.png')
    if os.path.exists(jpg):
        return jpg, jpeg_size(jpg)
    if os.path.exists(png):
        return png, png_size(png)
    return None, None


# ── المستند ─────────────────────────────────────────────────
doc = Document()

# اتجاه المقطع كلّه، وهوامش A4
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
sec.left_margin = sec.right_margin = Inches(0.75)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.8)
_set(sec._sectPr, 'w:bidi')

# الخطوط الافتراضية — ومعها خطّ النصوص المعقّدة (العربية) وإلا رجع Word إلى خطٍّ آخر
normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(11)
rpr = normal.element.get_or_add_rPr()
_set(rpr, 'w:rFonts', **{'w:ascii': FONT, 'w:hAnsi': FONT, 'w:cs': FONT})
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color in (('Heading 1', 20, ACCENT), ('Heading 2', 15, ACCENT),
                          ('Heading 3', 12.5, RGBColor(0x3D, 0x55, 0x48)),
                          ('Heading 4', 11.5, RGBColor(0x3D, 0x55, 0x48))):
    st = doc.styles[name]
    st.font.name = FONT
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    r = st.element.get_or_add_rPr()
    _set(r, 'w:rFonts', **{'w:ascii': FONT, 'w:hAnsi': FONT, 'w:cs': FONT})


def para(text='', style=None, align=None, rtl=True):
    p = doc.add_paragraph(style=style)
    if rtl:
        rtl_paragraph(p)
    if align is not None:
        p.alignment = align
    if text:
        add_inline(p, text, rtl)
    return p


# ── تنسيق داخل السطر: **عريض** و`شيفرة` و[نصّ](رابط) ────────
INLINE = re.compile(r'(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')


def add_inline(p, text, rtl=True):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = p.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = MONO
            r.font.size = Pt(9.5)
            rr = r._r.get_or_add_rPr()
            _set(rr, 'w:rFonts', **{'w:ascii': MONO, 'w:hAnsi': MONO, 'w:cs': MONO})
            rtl_run(r, False)
            continue
        elif part.startswith('[') and '](' in part:
            label = part[1:part.index('](')]
            r = p.add_run(label)
            r.font.color.rgb = ACCENT
            r.underline = True
        else:
            r = p.add_run(part)
        rtl_run(r, rtl)


def add_picture(name, caption):
    path, size = picture_path(name)
    if not path:
        return False
    w = MAX_W
    if size:
        px_w, px_h = size
        h = w * px_h / px_w
        if h > MAX_H:
            w = MAX_H * px_w / px_h
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(w))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rtl_paragraph(c)
        r = c.add_run(caption)
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
        r.italic = True
        rtl_run(r)
    return True


def add_table(head, rows):
    t = doc.add_table(rows=1, cols=len(head))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # الجدول يُقرأ من اليمين: العمود الأول يمينَ الصفحة
    _set(t._tbl.tblPr, 'w:bidiVisual')
    for i, h in enumerate(head):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        rtl_paragraph(p)
        add_inline(p, h)
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9.5)
        shade(cell._tc.get_or_add_tcPr(), 'EAF5EF')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row[:len(head)]):
            p = cells[i].paragraphs[0]
            rtl_paragraph(p)
            add_inline(p, val)
            for r in p.runs:
                r.font.size = Pt(9.5)
    doc.add_paragraph()


def add_code(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    shade(p._p.get_or_add_pPr(), 'F2F5F3')
    for i, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.name = MONO
        r.font.size = Pt(9)
        rr = r._r.get_or_add_rPr()
        _set(rr, 'w:rFonts', **{'w:ascii': MONO, 'w:hAnsi': MONO, 'w:cs': MONO})
        if i < len(lines) - 1:
            r.add_break()


def add_quote(lines):
    p = doc.add_paragraph()
    rtl_paragraph(p)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.15)
    shade(p._p.get_or_add_pPr(), 'FFF8E6')
    add_inline(p, ' '.join(lines))
    for r in p.runs:
        r.font.size = Pt(10.5)


# ── الغلاف ──────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(190)
r = t.add_run('دليل المستخدم')
r.font.size = Pt(34)
r.bold = True
r.font.color.rgb = ACCENT
rtl_run(r)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('منصة مركز تمكين الكفاءات')
r.font.size = Pt(15)
r.font.color.rgb = MUTED
rtl_run(r)

n = doc.add_paragraph()
n.alignment = WD_ALIGN_PARAGRAPH.CENTER
n.paragraph_format.space_before = Pt(26)
n.paragraph_format.left_indent = Inches(0.9)
n.paragraph_format.right_indent = Inches(0.9)
r = n.add_run('دليل تشغيلي لموظّفي المركز. يشرح لكل شاشة: مَن يفتحها، وكيف يعمل فيها '
              'خطوةً خطوة، وما الذي يمنعه النظام ولماذا.')
r.font.size = Pt(11)
r.font.color.rgb = MUTED
rtl_run(r)

d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
d.paragraph_format.space_before = Pt(22)
r = d.add_run(date.today().strftime('%Y-%m-%d'))
r.font.size = Pt(10.5)
r.font.color.rgb = MUTED

# ترقيم الصفحات في التذييل
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'PAGE')
footer._p.append(fld)

# ── فهرس Word حقيقي ────────────────────────────────────────
# حقلٌ يبنيه Word من العناوين بأرقام صفحاتها. وحتى لا يفتحه القارئ فيجده
# فارغاً، يُطلَب من Word تحديث الحقول عند الفتح (updateFields).
h = doc.add_heading(level=1)
rtl_paragraph(h)
h.paragraph_format.page_break_before = True
rtl_run(h.add_run('الفهرس'))

toc = doc.add_paragraph()
rtl_paragraph(toc)
run = toc.add_run()
_set(run._r, 'w:fldChar', **{'w:fldCharType': 'begin'})
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = r'TOC \o "1-2" \h \z \u'
run._r.append(instr)
_set(run._r, 'w:fldChar', **{'w:fldCharType': 'separate'})
hint = toc.add_run('يُولَّد الفهرس عند فتح الملف — وإن لم يظهر فاضغط بزرّ الفأرة '
                   'الأيمن هنا ثم «تحديث الحقل».')
hint.font.size = Pt(10)
hint.font.color.rgb = MUTED
rtl_run(hint)
_set(toc.add_run()._r, 'w:fldChar', **{'w:fldCharType': 'end'})

_set(doc.settings.element, 'w:updateFields', **{'w:val': 'true'})


# ── قراءة المصدر ────────────────────────────────────────────
md = open(SRC, encoding='utf-8').read()
md = re.sub(r'^#\s+.*\n', '', md, count=1)          # العنوان على الغلاف
md = re.sub(r'^##\s*الفهرس\s*\n[\s\S]*?(?=^##\s)', '', md, flags=re.M)  # فهرس Word يُولَّد
# تمهيد المستند يُقرأ على الغلاف لا بعد الفهرس، وإلا وقع بين الفهرس وأول فصل
md = re.sub(r'\A(?:>.*\n|\s*\n)+(?=---)', '', md)
md = re.sub(r'\A\s*---\s*\n', '', md)

lines = md.split('\n')
i = 0
img_ok = img_missing = 0
chapters = []

cells_of = lambda ln: [c.strip() for c in ln.strip().strip('|').split('|')]

while i < len(lines):
    ln = lines[i]

    if ln.startswith('```'):
        i += 1
        buf = []
        while i < len(lines) and not lines[i].startswith('```'):
            buf.append(lines[i]); i += 1
        i += 1
        add_code(buf)
        continue

    m = re.match(r'^(#{2,4})\s+(.*)$', ln)
    if m:
        level = len(m.group(1))
        title = re.sub(r'[*`]', '', m.group(2)).strip()
        h = doc.add_heading(level=level - 1)
        rtl_paragraph(h)
        if level == 2:
            # فاصل صفحة على العنوان نفسه لا بفقرةٍ فارغة قبله: الفقرة الفارغة
            # تبقى في المستند علامةً بيضاء يراها القارئ ويحاول حذفها.
            chapters.append(title)
            h.paragraph_format.page_break_before = True
        r = h.add_run(title)
        rtl_run(r)
        i += 1
        continue

    if re.match(r'^---+\s*$', ln):
        i += 1
        continue

    mi = re.match(r'^\s*!\[([^\]]*)\]\(guide-images/([^)]+)\.png\)\s*$', ln)
    if mi:
        if add_picture(mi.group(2), mi.group(1)):
            img_ok += 1
        else:
            img_missing += 1
        i += 1
        continue

    if ln.strip().startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i + 1]):
        head = cells_of(ln)
        i += 2
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            rows.append(cells_of(lines[i])); i += 1
        add_table(head, rows)
        continue

    if ln.startswith('>'):
        buf = []
        while i < len(lines) and lines[i].startswith('>'):
            buf.append(lines[i].lstrip('>').strip()); i += 1
        add_quote([b for b in buf if b])
        continue

    ml = re.match(r'^(\s*)([-*]|\d+[.)])\s+(.*)$', ln)
    if ml:
        # تُجمَع أسطر العنصر الواحد قبل تنسيقها: العريض قد يمتدّ على سطرين في
        # المصدر، وتنسيقُ كل سطر وحده كان يترك نجمتيه ظاهرتين في المستند.
        items = []
        while i < len(lines):
            m2 = re.match(r'^(\s*)([-*]|\d+[.)])\s+(.*)$', lines[i])
            if m2:
                items.append([m2.group(2)[0].isdigit(), m2.group(3)])
                i += 1
                continue
            if items and lines[i].strip() and re.match(r'^\s{2,}\S', lines[i]):
                items[-1][1] += ' ' + lines[i].strip()
                i += 1
                continue
            break
        for numbered, text in items:
            p = doc.add_paragraph(style='List Number' if numbered else 'List Bullet')
            rtl_paragraph(p)
            add_inline(p, text)
        continue

    if ln.strip() == '':
        i += 1
        continue

    buf = []
    while i < len(lines) and lines[i].strip() and not re.match(
            r'^(#{2,4}\s|```|>|\s*([-*]|\d+[.)])\s|---+\s*$)', lines[i]) and not lines[i].strip().startswith('|'):
        buf.append(lines[i].strip()); i += 1
    text = ' '.join(buf)
    inline_img = re.search(r'!\[([^\]]*)\]\(guide-images/([^)]+)\.png\)', text)
    if inline_img:
        text = text.replace(inline_img.group(0), '').strip()
        if text:
            para(text)
        if add_picture(inline_img.group(2), inline_img.group(1)):
            img_ok += 1
        else:
            img_missing += 1
        continue
    if text:
        para(text)

doc.save(OUT)
size = os.path.getsize(OUT) / 1048576
print(f'✓ {OUT} — {size:.1f} ميغابايت')
print(f'  {len(chapters)} فصلاً · {img_ok} صورة مضمَّنة'
      + (f' · {img_missing} مفقودة' if img_missing else ''))
